# -*- coding: utf-8 -*-
"""在课程模板基础上生成《Copilot 远程操控系统》期末报告正文。
- 保留模板：封面、成绩考核表、报告要求三部分及其原格式
- 仅填写封面“题目”，更新报告完成时间
- 追加：题目页(摘要/关键词) + 目录(TOC域) + 八章正文 + 参考文献
- 标题套用 Heading 1/2/3（中文字体），正文用宋体小四首行缩进2字符
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOC_DIR = os.path.dirname(HERE)
SRC = os.path.join(DOC_DIR, "template_converted.docx")
OUT = os.path.join(DOC_DIR, "Copilot远程操控系统-期末课程报告.docx")

doc = Document(SRC)

# ----------------------------------------------------------------------------
# 样式工具
# ----------------------------------------------------------------------------
def set_cjk(obj_font_elem, latin, ea):
    rpr = obj_font_elem
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), ea)


def ensure_para_style(name):
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles['Normal']
    return st


def config_style(name, latin, ea, size, bold, color=None, before=6, after=6,
                 line=1.5, line_rule=WD_LINE_SPACING.MULTIPLE, first_indent_chars=0,
                 align=None, outline=None):
    st = doc.styles[name]
    st.font.size = Pt(size)
    st.font.bold = bold
    if color:
        st.font.color.rgb = RGBColor(*color)
    rpr = st.element.get_or_add_rPr()
    set_cjk(rpr, latin, ea)
    pf = st.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = line_rule
    if align is not None:
        pf.alignment = align
    ppr = st.element.get_or_add_pPr()
    if first_indent_chars:
        ind = ppr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind'); ppr.append(ind)
        ind.set(qn('w:firstLineChars'), str(first_indent_chars * 100))
        ind.set(qn('w:firstLine'), str(int(size * 20 * first_indent_chars)))
    if outline is not None:
        ol = ppr.find(qn('w:outlineLvl'))
        if ol is None:
            ol = OxmlElement('w:outlineLvl'); ppr.append(ol)
        ol.set(qn('w:val'), str(outline))
    return st


# 正文样式（宋体小四，1.5 倍行距，首行缩进 2 字符）
config_style('Normal', '宋体', '宋体', 12, False, before=0, after=0, line=1.5)
BODY = 'RptBody'
ensure_para_style(BODY)
config_style(BODY, '宋体', '宋体', 12, False, before=2, after=2, line=1.5,
             first_indent_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
# 标题样式：黑体（自定义，带大纲级别，供目录域识别）
H1, H2, H3 = 'RptH1', 'RptH2', 'RptH3'
for nm in (H1, H2, H3):
    ensure_para_style(nm)
config_style(H1, '黑体', '黑体', 16, True, color=(0, 0, 0), before=14, after=8,
             line=1.5, align=WD_ALIGN_PARAGRAPH.LEFT, outline=0)
config_style(H2, '黑体', '黑体', 14, True, color=(0, 0, 0), before=10, after=6,
             line=1.5, align=WD_ALIGN_PARAGRAPH.LEFT, outline=1)
config_style(H3, '黑体', '黑体', 12, True, color=(0, 0, 0), before=8, after=4,
             line=1.5, align=WD_ALIGN_PARAGRAPH.LEFT, outline=2)

# ----------------------------------------------------------------------------
# 1. 填写封面题目 / 更新完成时间
# ----------------------------------------------------------------------------
for p in doc.paragraphs:
    t = p.text
    if t.replace(' ', '').startswith('题目') and (':' in t or '：' in t):
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = '题    目:\tCopilot 远程操控系统'
        else:
            p.add_run('题    目:\tCopilot 远程操控系统')
    if '报告完成时间' in t:
        for r in p.runs:
            if '2025' in r.text:
                r.text = r.text.replace('2025', '2026')
    if '电子邮件' in t or '邮件' in t:
        # 清除模板遗留的邮箱超链接，保持个人信息为空
        for hl in p._p.findall(qn('w:hyperlink')):
            p._p.remove(hl)
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = '电子邮件:\t'
        else:
            p.add_run('电子邮件:\t')

# ----------------------------------------------------------------------------
# 2. 追加内容的辅助函数
# ----------------------------------------------------------------------------
def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    run._r.append(br)


def h1(text): doc.add_paragraph(text, style=H1)
def h2(text): doc.add_paragraph(text, style=H2)
def h3(text): doc.add_paragraph(text, style=H3)


def p(text, indent=True):
    para = doc.add_paragraph(style=BODY if indent else 'Normal')
    if not indent:
        para.paragraph_format.line_spacing = 1.5
    para.add_run(text)
    return para


def bullet(text):
    para = doc.add_paragraph(style=BODY)
    para.paragraph_format.first_line_indent = Pt(0)
    pf = para.paragraph_format
    pf.left_indent = Cm(0.74)
    para.add_run('· ' + text)
    return para


def caption(text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(10)
    r = para.add_run(text)
    r.font.size = Pt(10.5)
    set_cjk(r._r.get_or_add_rPr(), '宋体', '宋体')
    return para


def figure(filename, cap, width=14.6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(os.path.join(HERE, filename), width=Cm(width))
    caption(cap)


def shade(cell, color):
    # 表格不再着色，注释掉填充逻辑
    pass
    # tcpr = cell._tc.get_or_add_tcPr()
    # sh = OxmlElement('w:shd')
    # sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), color)
    # tcpr.append(sh)


def set_cell(cell, text, bold=False, size=10.5, header=False, align='left'):
    cell.text = ''
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.15
    para.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            para = cell.add_paragraph()
            para.paragraph_format.space_before = Pt(1); para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.15
            para.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
        r = para.add_run(line)
        r.font.size = Pt(size)
        r.font.bold = bold or header
        set_cjk(r._r.get_or_add_rPr(), '宋体', '宋体')
        # 表头不再着色，仅保持黑色粗体
        # if header:
        #     r.font.color.rgb = RGBColor(0x1c, 0x27, 0x33)
    # if header:
    #     shade(cell, 'D9E2F3')


def table(headers, rows, widths=None, cap=None, align_cols=None, size=10.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    aligns = align_cols or ['left'] * len(headers)
    for j, hd in enumerate(headers):
        set_cell(t.rows[0].cells[j], hd, header=True, align='center', size=size)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], str(val), align=aligns[j], size=size)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if cap:
        caption(cap)
    return t


def toc_field():
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    hint = OxmlElement('w:t'); hint.text = '右键“更新域”可生成目录'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep)
    run._r.append(hint); run._r.append(fld_end)


def centered_title(text, size, bold=True, ea='黑体', space_after=6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.5
    r = para.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold
    set_cjk(r._r.get_or_add_rPr(), ea, ea)
    return para


# ============================================================================
#  题目页：题目 / 作者 / 摘要 / 关键词
#  （模板末段带”下一页”分节符，已自动换页，无需再插入分页）
# ============================================================================
centered_title('Copilot 远程操控系统', 20, ea='黑体', space_after=2)
centered_title('——面向本地 AI 编码代理的安全远程操控与流式交互系统', 13, ea='黑体', space_after=14)

aut = doc.add_paragraph(); aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = aut.add_run('作者：______________（学号：____________  专业：网络空间安全）')
ar.font.size = Pt(12); set_cjk(ar._r.get_or_add_rPr(), '宋体', '宋体')
doc.add_paragraph()

centered_title('内容摘要', 14)
abstract = (
    '随着大语言模型驱动的 AI 编码代理（如 GitHub Copilot CLI）逐步具备在本地自主执行命令、'
    '读写文件、调用工具链的能力，其本质已从”代码补全工具”演进为一类可执行任意操作的本地自治代理。'
    '如何在保证安全的前提下，将这类强能力的本地代理安全地暴露给远程用户、实现移动端随时随地的远程操控，'
    '成为一个兼具工程价值与网络空间安全意义的问题。本项目设计并实现了”Copilot 远程操控系统”，'
    '以本机已登录的 Copilot CLI 为执行内核，通过一个通道无关的即时通讯（IM）机器人接入层，'
    '将本地 AI 代理桥接到移动端，实现远程指令下发、流式过程回传、工具执行可视化与多会话管理。'
    '系统把”远程操控本地代理”视为一条需要严格管控的远程命令执行通道，并以安全设计为核心：'
    '在身份层采用 chat_id / open_id 白名单统一鉴权；在通信层管理机器人凭据、支持代理与长连接/长轮询；'
    '在执行层通过工具授权、路径白名单与工作区隔离约束代理的操作边界；在会话层对会话标识做严格的路径校验，'
    '杜绝路径穿越；在运行层提供超时、降级与看门狗保活，保证长时间无人值守下的稳定性。'
    '系统采用 Python 异步架构实现，对 Copilot CLI 的 JSON 流式事件进行实时解析与中文摘要，'
    '并以流式文本、交互卡片、工具执行摘要与状态诊断报告等多种形式呈现数据。'
    '测试表明，系统能够稳定地完成远程对话、流式展示、会话接管与异常恢复，'
    '在可用性与安全可控性之间取得了较好平衡。'
)
p(abstract)
kp = doc.add_paragraph(style=BODY)
kr = kp.add_run('关键词：')
kr.font.bold = True; set_cjk(kr._r.get_or_add_rPr(), '黑体', '黑体')
kp.add_run('Copilot CLI；AI 编码代理；远程操控；即时通讯机器人；访问控制；流式交互；会话管理；网络空间安全')

# Abstract 与 Keywords 已删除（非正式论文，不需要英文摘要）

# ============================================================================
#  目录
# ============================================================================
add_page_break()
centered_title('目  录', 16)
toc_field()

# ============================================================================
#  正文
# ============================================================================
add_page_break()

# ---------- 一、选题背景与意义 ----------
h1('一、选题背景与意义')
h2('1.1 选题背景')
p('近年来，以大语言模型（LLM）为核心的 AI 编码助手快速普及。GitHub Copilot 等产品已经从'
  '编辑器内的“代码补全”，发展出独立的命令行代理形态——Copilot CLI。这类命令行 AI 代理不再只是'
  '给出文本建议，而是能够在本地主机上自主地执行 Shell 命令、读写文件、检索代码、调用各类工具，'
  '并以多轮对话的方式持续推进一项工程任务。换言之，它已成为一个运行在本机、具备真实操作能力的'
  '自治代理（Autonomous Agent）。')
p('与此同时，开发者的工作场景越来越碎片化与移动化：长耗时的构建、测试、代码分析任务往往在'
  '一台固定的工作主机上运行，而开发者却可能临时离开工位，仅随身携带移动设备。如何在离开主机后，'
  '依然能够安全地查看任务进展、追加指令、接管正在运行的会话，是一个普遍而现实的需求。')
p('本项目正是在这一背景下提出：以本机已登录的 Copilot CLI 作为执行内核，构建一套“Copilot 远程'
  '操控系统”，让用户能够通过随身的即时通讯（IM）客户端，远程操控这台主机上的 AI 代理，'
  '实现指令下发、流式过程查看、会话接管与多会话切换。')
h2('1.2 现实问题与痛点')
p('要把一个具备本地执行能力的 AI 代理“搬到”远程移动端，并非简单的消息转发，核心难点集中在以下方面：')
bullet('能力与风险并存：Copilot CLI 能执行命令、改写文件，一旦远程通道被滥用，等价于把本机的'
       '操作权限暴露给攻击者，因此远程操控本质上是一条必须严格管控的远程命令执行通道。')
bullet('上下文连续性：AI 代理是有状态、多轮的，远程端需要能够准确绑定、恢复并接管本机已有会话，'
       '而不是每次都从零开始。')
bullet('过程可见性：代理在执行过程中会产生大量中间事件（思考、工具调用、命令输出），'
       '移动端屏幕有限，需要实时、可读地呈现“流式过程”，而非等待数分钟后才返回一坨结果。')
bullet('无人值守的稳定性：远程操控通常长时间运行，网络抖动、代理中断、进程异常都可能导致'
       '“进程还在、却不再消费消息”的假死，需要恢复与保活机制。')
h2('1.3 选题意义')
p('本选题归属于网络空间安全实训中“网络编程 / 通信安全 / 系统安全”方向，其安全意义体现在：'
  '本系统是一个典型的“远程控制 + 本地命令执行”场景，与网络空间安全中经常讨论的远程访问、'
  '最小权限、纵深防御等议题高度契合。项目并非简单实现“能用”的功能，而是把安全可控性作为第一性目标，'
  '系统化地设计了身份认证、通信安全、执行边界、会话边界与运行保障五道防线（详见第五章）。')
p('从工程意义看，系统验证了“将本地强能力 AI 代理安全外延到移动端”这一范式的可行性，'
  '为后续研究 AI 代理的远程治理、审计与权限收敛提供了一个可运行的参考实现。')

# ---------- 二、文献综述 ----------
h1('二、文献综述')
h2('2.1 AI 编码代理与命令行 Agent 的发展')
p('AI 编码工具的演进大致经历了三个阶段：第一阶段是编辑器内的补全与建议；第二阶段是对话式'
  '编程助手，能够理解自然语言意图并生成代码片段；第三阶段则是具备工具使用（Tool Use）能力的'
  '自治代理，能够自主拆解任务、调用命令行工具、读写工作区文件并循环迭代。Copilot CLI 属于第三阶段，'
  '其与传统补全工具的根本区别在于“可执行性”——它在本地拥有真实的操作能力。这一能力既是其价值所在，'
  '也是其安全风险的来源。')
h2('2.2 即时通讯机器人与远程接入')
p('即时通讯（IM）平台普遍提供机器人（Bot）开放能力，常见接入方式有两类：一类是基于 HTTP 长轮询'
  '（Long Polling）的公有 Bot 平台，客户端周期性拉取更新；另一类是基于长连接（WebSocket）的'
  '事件推送，如飞书开放平台的 lark-oapi 长连接客户端。IM 机器人天然具备身份体系、消息可达性与'
  '富交互（卡片、菜单、按钮）能力，非常适合作为“远程操控”的人机接口。本项目即利用这一点，'
  '将 IM Bot 作为远程操控本地 AI 代理的统一入口，并以飞书长连接通道为主、辅以一个基于长轮询的'
  '公有 IM Bot 通道，形成通道无关的接入层。')
h2('2.3 远程命令执行的安全研究现状')
p('远程命令执行（RCE）一直是网络空间安全的核心议题。无论是合法的远程运维，还是恶意的远程控制，'
  '其安全模型都围绕“谁可以执行、能执行什么、在哪里执行、出了问题如何止损”展开，对应身份认证、'
  '授权与最小权限、执行环境隔离、审计与熔断等机制。把 AI 代理放到远程操控场景下，这些经典安全原则'
  '同样适用，且更为关键——因为代理具有一定的自主性，可能在一次指令下触发一连串工具调用与文件操作。'
  '因此本项目在设计上充分借鉴了访问控制白名单、最小权限授权、路径隔离与防路径穿越、超时熔断等'
  '安全工程实践。')
h2('2.4 现有方案对比与本项目定位')
table(
    ['方案类型', '远程能力', '上下文/会话', '安全控制', '局限'],
    [
        ['SSH / 远程桌面', '强（完整主机）', '不针对 AI 代理', '账号+密钥', '权限过大，移动端体验差，难以对单个 AI 会话治理'],
        ['云端 AI 编程平台', '强', '云端会话', '平台托管', '需上云、脱离本机环境与已登录凭据'],
        ['普通 IM 通知机器人', '弱（仅通知）', '无', '基本无', '只能单向推送，无法操控与接管'],
        ['本项目', '聚焦 AI 代理', '绑定/接管本机会话', '五道安全防线', '依赖本机常驻，需自建机器人'],
    ],
    widths=[2.6, 2.4, 2.6, 2.4, 4.2],
    align_cols=['center', 'center', 'center', 'center', 'left'],
    cap='表 2-1 远程操控相关方案对比',
    size=10,
)
p('综上，本项目定位为“面向本地 AI 编码代理的、以安全为核心的远程操控系统”，'
  '在保留本机环境与已登录凭据的同时，提供受控、可审计、可接管的远程操控能力，'
  '填补了上述方案在“AI 代理细粒度远程治理”上的空白。')

# ---------- 三、研究（开发）的目的 ----------
h1('三、研究（开发）的目的')
h2('3.1 总体目标')
p('构建一套可在本机常驻运行的 Copilot 远程操控系统，使授权用户能够通过 IM 客户端安全地远程操控'
  '本机 Copilot CLI，完成多轮对话、流式过程查看、会话管理与接管，并在通信、执行与运行各层面'
  '提供可验证的安全控制。')
h2('3.2 具体目标')
p('功能性目标：')
bullet('实现通道无关的 IM 接入层，至少支持飞书长连接与一个长轮询公有 IM Bot 两种通道。')
bullet('实现远程对话：用户发送文本即进入当前会话，系统调用 Copilot CLI 并实时回传过程与结果。')
bullet('实现多会话管理：新建、列表、按前缀切换、接管本机运行中会话、查看历史、删除。')
bullet('实现模型查看与切换，并能从 CLI 实时发现可用模型。')
p('安全性目标：')
bullet('实现基于 chat_id / open_id 白名单的统一访问控制。')
bullet('对会话标识实施严格路径校验，杜绝绝对路径与“../”路径穿越。')
bullet('提供工具授权、路径白名单与工作区隔离等执行边界控制。')
p('非功能性目标：流式低延迟呈现、长时间无人值守的稳定性（超时、降级、看门狗保活）、'
  '跨平台（Windows PowerShell 环境优先）与良好的可诊断性（doctor 启动前体检）。')

# ---------- 四、需求分析 ----------
h1('四、需求分析')
h2('4.1 总体需求与系统角色')
p('系统涉及两类角色：一是“远程用户”，通过 IM 客户端下发指令、查看过程与结果、管理会话；'
  '二是“本机运维者”，在主机命令行完成初始化、体检与启动。系统的核心用例包括：远程对话、'
  '流式过程查看、会话管理与接管、模型切换、状态诊断。')
h2('4.2 功能需求')
table(
    ['编号', '功能', '说明', '优先级'],
    [
        ['FR-1', '远程对话', '接收文本/富文本消息，进入当前会话上下文并调用 Copilot CLI', '高'],
        ['FR-2', '流式过程回传', '实时回传思考、工具调用与命令输出摘要', '高'],
        ['FR-3', '流式结果呈现', '以消息编辑或分段方式增量呈现最终回复', '高'],
        ['FR-4', '会话新建/切换', '新建会话、按唯一前缀切换已保存会话', '高'],
        ['FR-5', '会话接管', '发现并接管本机运行中的 Copilot 会话，延续上下文', '高'],
        ['FR-6', '会话历史/删除', '查看会话历史摘要、刷新追踪、删除会话', '中'],
        ['FR-7', '模型查看/切换', '查看当前模型，实时发现并切换可用模型', '中'],
        ['FR-8', '访问控制', '基于白名单对所有入口统一鉴权', '高'],
        ['FR-9', '状态诊断', '查看后端状态与诊断报告（/status、doctor）', '中'],
    ],
    widths=[1.4, 2.2, 8.2, 1.4],
    align_cols=['center', 'center', 'left', 'center'],
    cap='表 4-1 功能需求列表',
    size=10,
)
h2('4.3 非功能需求')
table(
    ['类别', '需求描述', '指标/约束'],
    [
        ['性能', '流式过程应低延迟呈现，避免长时间无反馈', '过程消息编辑限频；增量回传'],
        ['安全', '仅授权身份可访问；约束代理操作边界', '白名单鉴权；路径白名单；防路径穿越'],
        ['可用性', '长时间无人值守稳定运行，异常可恢复', '单次调用超时 3600s；断线重连；看门狗保活'],
        ['兼容性', '兼容 Windows PowerShell 与 .bat/.cmd/.ps1 命令入口', 'argv 转义；ps1 经 powershell 启动'],
        ['可维护性', '可在启动前体检并定位问题', 'doctor 输出配置/目录/命令/工作区诊断'],
    ],
    widths=[1.8, 6.6, 5.0],
    align_cols=['center', 'left', 'left'],
    cap='表 4-2 非功能需求',
    size=10,
)
h2('4.4 安全需求与威胁分析')
p('由于远程操控等价于赋予远端对本机 AI 代理的指挥权，必须正视其威胁面。结合 STRIDE 思路，'
  '主要威胁与对应需求如下：')
table(
    ['威胁', '场景', '对应安全需求'],
    [
        ['假冒/越权访问', '非授权用户向机器人发送指令', '身份白名单鉴权，未授权直接拒绝'],
        ['凭据泄露', '机器人 Token / app_secret 被窃取', '凭据本地存储、最小暴露、可配代理'],
        ['权限滥用', '通过代理执行越界命令或访问越界路径', '工具授权、路径白名单、工作区隔离'],
        ['路径穿越', '构造 ../ 或绝对路径读取/删除任意会话', '会话标识仅限 session-state 直接子目录'],
        ['拒绝服务/假死', '长耗时任务或网络抖动导致进程假死', '调用超时、降级提示、断线重建与保活'],
    ],
    widths=[2.2, 5.6, 5.6],
    align_cols=['center', 'left', 'left'],
    cap='表 4-3 威胁分析与安全需求映射',
    size=10,
)

# ---------- 五、总体设计与详细设计 ----------
h1('五、总体设计与详细设计')
h2('5.1 总体架构')
p('系统采用分层架构，自上而下分为终端接入层、IM 通道接入层、业务编排层与 Copilot 执行层，'
  '并辅以配置诊断、持久化、会话扫描与运行保障等支撑模块。其核心思想是：把“与具体 IM 平台相关的'
  '协议细节”收敛在通道接入层，业务编排层只面向统一的抽象，从而做到通道无关、易于扩展。'
  '总体架构如图 5-1 所示。')
figure('fig_arch.png', '图 5-1 系统总体架构')
p('图中，终端接入层是用户与运维者的入口；IM 通道接入层负责协议适配与统一的白名单鉴权；'
  '业务编排层以 TaskRunner 为中心，汇总上下文并通过 LiveProgress 流式协议与通道解耦；'
  'Copilot 执行层由 AssistantPlanner 组装命令、拉起 Copilot CLI 子进程并解析其 JSONL 事件，'
  '最终作用于本地工作区。')
h2('5.2 模块划分与职责')
p('系统按职责划分为若干高内聚、低耦合的模块，模块依赖关系如图 5-2 所示，主要职责见表 5-1。')
figure('fig_module.png', '图 5-2 模块结构与依赖关系')
table(
    ['模块', '代码位置', '主要职责'],
    [
        ['CLI 入口', 'cli/main.py', '提供 init/start/doctor，准备环境并分发启动'],
        ['配置与路径', 'config.py / paths.py', '目录结构、默认配置、JSON 读写与强校验、启动前诊断'],
        ['通道接入', 'telegram_bot.py / feishu_bot.py', '协议适配、白名单鉴权、过程/结果消息、菜单与卡片交互'],
        ['任务编排', 'task_runner.py', '汇总上下文、调用规划器、持久化、会话菜单与接管'],
        ['Copilot 调用', 'agent.py', '组装 argv、拉起子进程、解析 JSONL、工具摘要、CLI 诊断'],
        ['会话扫描', 'copilot_sessions.py', '扫描 session-state，解析工作区/模型/状态/历史'],
        ['持久化', 'session_store.py / conversation_store.py', '保存会话列表/当前会话/模型与最近对话历史'],
    ],
    widths=[1.8, 3.8, 7.6],
    align_cols=['center', 'left', 'left'],
    cap='表 5-1 模块职责一览',
    size=10,
)
h2('5.3 关键流程：单轮对话的流式时序')
p('远程一次对话从用户发送文本开始，经鉴权、上下文汇总、子进程调用，到逐行消费 JSONL 事件并'
  '流式回传，最终展示完整回复并持久化历史。完整时序如图 5-3 所示。')
figure('fig_seq.png', '图 5-3 单轮对话流式交互时序')
p('该流程的关键在于“边执行、边呈现”：系统不等待 Copilot CLI 全部结束，而是逐行读取其标准输出，'
  '将工具调用类事件转为过程日志、将 assistant.message_delta 增量转为流式回复，'
  '从而在移动端获得近实时的过程可见性。')
h2('5.4 详细设计')
h3('5.4.1 IM 通道接入与访问控制')
p('通道接入层为每个平台提供独立实现，但共享统一的鉴权与业务调用链。飞书侧采用 lark-oapi 长连接'
  '客户端，接收 im.message.receive_v1 文本/富文本消息、机器人菜单事件与卡片动作回调，菜单与快捷'
  '入口使用交互卡片，普通对话过程与结果使用文本分段发送；长轮询通道侧构建应用、注册命令、文本'
  '消息与按钮回调。两者均通过 restricted() 装饰器统一做白名单校验，未授权请求直接返回拒绝提示，'
  '避免鉴权逻辑散落在各 handler 中。')
h3('5.4.2 任务编排 TaskRunner 与流式协议 LiveProgress')
p('TaskRunner 是业务协调中心，负责汇总对话历史、当前会话、当前模型与工作区等上下文，调用'
  'AssistantPlanner 执行单轮对话，并在请求前后完成持久化写入；它不直接关心任何具体 IM 平台的'
  'API 细节，过程输出通过 LiveProgress 协议抽象，从而与具体消息实现解耦。会话键统一以字符串'
  '持久化，既兼容数字型 chat_id，也兼容字符串型 chat_id。')
h3('5.4.3 Copilot CLI 调用与 JSONL 事件解析')
p('AssistantPlanner 通过 --session-id 绑定/恢复会话，使用 --output-format json 获取结构化事件流，'
  '命令形如：copilot --session-id <id> --output-format json -p <prompt> -s；只有用户显式选择具体模型时才附带 --model。'
  '系统以异步子进程方式启动 Copilot CLI，并发消费 stdout/stderr：逐行尝试 JSON 解析，'
  '按事件类型分发——tool.execution_start / tool.execution_complete 归为过程日志并压缩为中文摘要，'
  'assistant.message_delta 归为流式回复增量，assistant.reasoning 归为“思考”摘要。'
  '对非零退出码、超时与空输出分别生成稳定的中文降级提示，并建议执行 doctor 排查。')
h3('5.4.4 会话扫描与接管')
p('系统直接复用 Copilot 原生状态目录 ~/.copilot/session-state，CopilotSessionInspector 解析其中的'
  'workspace.yaml（工作区、摘要）、events.jsonl（模型、最后事件时间、历史摘要）与 inuse.*.lock'
  '（是否运行中），据此实现远程的会话接管、历史预览与删除，机制如图 5-4 所示。这种“扫描复用”'
  '的方式无需额外服务接口，且能让远程端无缝延续本机已有上下文。')
figure('fig_session.png', '图 5-4 会话状态扫描与接管机制', width=13.5)
h3('5.4.5 持久化设计')
p('持久化模块以 JSON 文件保存每个 Chat 的最近对话历史（默认仅保留最近 40 条，防止无限增长）、'
  '当前会话、会话列表与当前模型。读取时跳过结构异常的记录，避免单条坏数据导致整个机器人无法启动；'
  '写入前确保父目录存在；当前激活会话必须存在于会话列表中，陈旧 active 值按无激活处理。')
h3('5.4.6 配置与诊断')
p('配置固定存放于 ~/.topilot/config.json，使用 Settings 数据类集中承载，解析函数统一转换布尔、'
  '整数、浮点与模型列表；白名单字段支持数组或逗号分隔字符串，非法项被忽略而不致整体加载失败。'
  'doctor 子命令在不启动通道、不真实调用 Copilot 的前提下，输出配置/目录/命令可执行性/工作区'
  '与问题清单，作为启动前体检。')
h2('5.5 安全设计（核心）')
p('安全是本系统的第一性目标。围绕“远程操控等价于远程命令执行”这一判断，系统构建了从入口到执行的'
  '五道纵深防线，如图 5-5 所示。')
figure('fig_security.png', '图 5-5 访问控制与权限边界（五道防线）')
p('① 身份认证：所有入口经 restricted() 统一拦截，仅放行白名单内的 chat_id / open_id，'
  '未授权请求直接拒绝，从源头阻断假冒与越权。')
p('② 通信安全：机器人 Token 与飞书 app_secret 仅存于本机配置，支持为长轮询通道配置代理；'
  '飞书侧使用加密长连接接收事件，减少凭据与数据的暴露面。')
p('③ 执行边界：通过 allow_all_tools 控制是否放开工具授权，通过 allow_all_paths 与 --add-dir'
  '控制代理可访问的路径范围（默认不放开全路径，仅以工作区与显式目录构成白名单），'
  '并以工作区为子进程 cwd 实现操作隔离，遵循最小权限原则。')
p('④ 会话边界：会话标识仅接受 session-state 根目录下的“直接子目录名”，拒绝绝对路径与含'
  '“../”的路径型标识，从代码层面杜绝借会话读取/删除越界目录的路径穿越攻击。')
p('⑤ 运行保障：单次调用设置超时（默认 3600 秒）并在超时后杀死子进程；对异常输出做稳定降级；'
  '配合看门狗与断线重建，避免长时间运行中的假死与失联。')
p('需要指出，allow_all_tools 与 allow_all_paths 属于“便利性与安全性”的权衡开关：在个人可信'
  '环境下可适度放开以提升效率，但在更严格的场景应保持收敛。系统通过默认值（allow_all_paths=False）'
  '与显式白名单，将默认姿态偏向安全。')
h2('5.6 接口与配置设计')
p('系统对外接口包括本机 CLI 命令与 IM 端命令两类，关键配置项见表 5-3。')
table(
    ['命令', '作用'],
    [
        ['topilot init [--force]', '交互式生成/覆盖配置文件'],
        ['topilot doctor', '启动前体检：配置/目录/命令/工作区/问题清单'],
        ['topilot start', '启动长轮询通道与/或飞书长连接'],
        ['/sessions、/session_new、/session_use', '会话列表、接管、新建、按前缀切换'],
        ['/model、/status、/whoami', '模型切换、后端状态诊断、身份查看'],
    ],
    widths=[5.4, 7.8],
    align_cols=['left', 'left'],
    cap='表 5-2 主要命令接口',
    size=10,
)
table(
    ['配置项', '默认值', '说明'],
    [
        ['copilot.model', 'auto', 'Copilot CLI 默认模型，auto 表示不传 --model'],
        ['copilot.timeout_seconds', '3600', '单次调用超时秒数'],
        ['copilot.allow_all_tools', 'true', '是否附带 --allow-all-tools'],
        ['copilot.allow_all_paths', 'false', '是否放开全路径访问（默认收敛）'],
        ['copilot.add_workspace_dir', 'true', '是否自动附带 --add-dir 工作区'],
        ['runtime.session_watch_interval_seconds', '2.0', '会话追踪轮询间隔（1–15s 生效）'],
        ['*.allowed_chat_ids / allowed_open_ids', '空', '访问白名单，留空风险高，强烈建议配置'],
    ],
    widths=[5.4, 2.4, 5.4],
    align_cols=['left', 'center', 'left'],
    cap='表 5-3 关键配置项',
    size=10,
)

# ---------- 六、系统实现与测试 ----------
h1('六、系统实现与测试')
h2('6.1 开发环境与技术栈')
table(
    ['项目', '内容'],
    [
        ['编程语言', 'Python 3.11+'],
        ['核心范式', 'asyncio 异步并发（子进程消费、并发收发）'],
        ['执行内核', 'GitHub Copilot CLI（需已 copilot login）'],
        ['飞书接入', 'lark-oapi 长连接客户端'],
        ['长轮询通道', '基于 long-polling 的公有 IM Bot SDK'],
        ['运行环境', 'Windows PowerShell（兼容 .bat/.cmd/.ps1 入口）'],
        ['打包/测试', 'pyproject + pytest'],
    ],
    widths=[3.2, 10.0],
    align_cols=['center', 'left'],
    cap='表 6-1 开发环境与技术栈',
    size=10,
)
h2('6.2 关键实现')
h3('6.2.1 流式输出与限频')
p('系统逐行读取 Copilot CLI 的 JSONL 标准输出，将 assistant.message_delta 的增量文本通过'
  'reply_streamer 回传，过程类事件通过 progress_logger 回传。为兼顾移动端可读性与平台接口稳定性，'
  '过程消息采用限频编辑/分段发送：飞书侧普通对话过程与最终回复以文本分段发送，菜单与快捷入口用卡片；'
  '长轮询通道侧则以消息编辑实现增量刷新。')
h3('6.2.2 工具调用过程的中文摘要')
p('代理执行过程中会触发 view/read_file、glob/list_dir、命令执行等多类工具事件，原始 payload 冗长'
  '且为英文。系统针对不同工具类型采用不同摘要策略，例如读取文件/目录、列目录、执行命令分别生成'
  '“读取文件：<path>”“执行命令”等简洁中文提示，并对过大的命令输出做截断或落盘提示，'
  '显著提升小屏可读性；同时跳过 report_intent 等纯规划类噪声事件。')
h3('6.2.3 断线恢复与看门狗')
p('长轮询通道为普通请求与 getUpdates 分别设置超时；当轮询层出现网络异常且无具体更新上下文时，'
  '错误处理器会标记需要重建并停止当前 Application，启动层在轮询返回后检查该标记并在等待后重建，'
  '从而缓解“代理短暂中断后进程仍在却不再消费更新”的假死。此外提供 PowerShell 看门狗脚本，'
  '按固定间隔巡检 topilot 进程并在缺失时自动拉起，保障无人值守运行。机制如图 6-1 所示。')
figure('fig_deploy.png', '图 6-1 部署形态与运行时保障', width=14.0)
h3('6.2.4 跨平台与 PowerShell 适配')
p('针对 Windows 下 .ps1 入口，系统通过 powershell -NoProfile -ExecutionPolicy Bypass -File 启动；'
  '对 .bat/.cmd 命令做换行转义，避免长 prompt 导致参数截断。模型列表优先从 CLI 主帮助的 --model'
  'choices 实时解析；默认 auto 且未配置候选列表时只展示 auto，其余场景失败后回退解析 help config，再失败则回退到配置列表，保证 /model 始终可用。')
h2('6.3 部署与运行')
p('部署流程简洁：在项目根目录执行 pip install -e . 安装，首次运行 topilot 自动进入交互式配置向导，'
  '生成默认配置 ~/.topilot/config.json；随后可依次执行 topilot init、topilot doctor、topilot start。'
  '当仅启用飞书时以 Feishu-only 模式常驻；两者都启用时，飞书长连接在后台线程启动、长轮询通道'
  '保持主线程运行。')
h2('6.4 测试与结果分析')
p('系统提供 pytest 单元测试，覆盖配置解析、持久化存储、会话扫描、Copilot 调用辅助与通道辅助'
  '逻辑等关键模块。除单元测试外，针对核心用例设计了功能测试，主要用例与结果见表 6-2。')
table(
    ['用例', '步骤', '预期结果', '结果'],
    [
        ['TC-1 白名单鉴权', '非白名单身份发送指令', '被拒绝并提示未授权', '通过'],
        ['TC-2 远程对话与流式', '发送一条任务指令', '过程日志与回复增量实时呈现', '通过'],
        ['TC-3 会话接管', '接管本机运行中的会话', '延续上下文，摘要与模型正确', '通过'],
        ['TC-4 路径穿越防护', '以 ../ 或绝对路径作会话标识', '被拒绝，不读取/删除越界目录', '通过'],
        ['TC-5 超时降级', '构造超时调用', '杀死子进程并返回稳定降级提示', '通过'],
        ['TC-6 断线恢复', '模拟轮询网络异常', '标记并重建通道，进程不退出', '通过'],
        ['TC-7 模型发现切换', '执行 /model 切换模型', '实时列出并成功切换', '通过'],
    ],
    widths=[2.6, 4.0, 4.6, 1.4],
    align_cols=['center', 'left', 'left', 'center'],
    cap='表 6-2 主要功能测试用例与结果',
    size=10,
)
p('结果分析：在白名单鉴权、路径穿越防护、超时降级与断线恢复等安全/稳定性用例上，系统行为均符合'
  '设计预期，验证了五道防线的有效性；在远程对话与流式呈现上，过程与结果能够近实时回传，移动端'
  '体验良好。整体表明系统在“强能力 + 可控性”之间达到了预期平衡。')

# ---------- 七、总结 ----------
h1('七、总结')
h2('7.1 项目核心内容与使用价值')
p('本项目实现了一套以本地 Copilot CLI 为执行内核、以通道无关 IM 接入层为人机接口的远程操控系统，'
  '使授权用户能够随时随地安全地操控本机 AI 编码代理。其使用价值在于：在不脱离本机环境与已登录'
  '凭据的前提下，提供受控、可接管、可审计的远程 AI 能力，适用于长耗时任务的远程跟进与移动办公。')
h2('7.2 特点与优点')
bullet('以安全为核心：把远程操控视为远程命令执行通道，系统化构建五道纵深防线。')
bullet('通道无关：业务编排与具体 IM 平台解耦，易于扩展新通道。')
bullet('上下文连续：直接复用 Copilot 原生 session-state，实现无缝接管与历史延续。')
bullet('体验友好：流式过程、中文摘要与富交互兼顾了移动端可读性。')
bullet('稳定可诊断：超时、降级、断线重建与看门狗保障无人值守运行，doctor 支持启动前体检。')
h2('7.3 不足与改进方向')
bullet('审计与追溯：当前以日志为主，后续可引入结构化操作审计与命令级回放。')
bullet('权限粒度：可进一步细化到“按用户/按会话”的工具与路径授权，替代全局开关。')
bullet('凭据保护：可引入操作系统密钥库或加密存储，替代明文配置文件。')
bullet('多用户与并发：可完善并发会话隔离与配额，支持团队场景。')
bullet('对抗性测试：后续可补充针对提示注入、越权诱导等 AI 特有攻击面的专项测试。')
p('总体而言，本项目较好地完成了选题目标，验证了“安全地把本地 AI 代理外延到移动端”的可行性，'
  '并在网络空间安全的视角下，对远程命令执行场景的访问控制、最小权限、隔离与熔断等机制做了'
  '完整的工程落地，具有进一步研究与扩展的价值。')

# ---------- 八、主要参考文献 ----------
h1('八、主要参考文献')
refs = [
    '[1] GitHub. GitHub Copilot Documentation[EB/OL]. https://docs.github.com/copilot, 2024.',
    '[2] Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020.',
    '[3] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]. International Conference on Learning Representations (ICLR), 2023.',
    '[4] Schick T, Dwivedi-Yu J, Dessì R, et al. Toolformer: Language Models Can Teach Themselves to Use Tools[C]. NeurIPS, 2023.',
    '[5] Saltzer J H, Schroeder M D. The Protection of Information in Computer Systems[J]. Proceedings of the IEEE, 1975, 63(9): 1278-1308.',
    '[6] Shostack A. Threat Modeling: Designing for Security[M]. Wiley, 2014.',
    '[7] 飞书开放平台. 服务端 API 与长连接（lark-oapi）开发文档[EB/OL]. https://open.feishu.cn/document, 2024.',
    '[8] Python Software Foundation. asyncio — Asynchronous I/O[EB/OL]. https://docs.python.org/3/library/asyncio.html, 2024.',
    '[9] OWASP Foundation. OWASP Top 10[EB/OL]. https://owasp.org/Top10, 2021.',
    '[10] Hunt R. Network security: the principles of threats, attacks and intrusions[J]. Computer Communications, 2004.',
]
for r in refs:
    para = doc.add_paragraph(style='Normal')
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.left_indent = Cm(0.85)
    para.paragraph_format.first_line_indent = Cm(-0.85)
    run = para.add_run(r)
    run.font.size = Pt(10.5)
    set_cjk(run._r.get_or_add_rPr(), '宋体', '宋体')

doc.save(OUT)
print('SAVED', OUT)
